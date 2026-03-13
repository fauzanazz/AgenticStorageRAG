"""Tests for document processors - DOCX."""

from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument

from app.domain.documents.processors.docx import DocxProcessor


def _create_test_docx(paragraphs: list[str], headings: list[str] | None = None) -> bytes:
    """Create a simple DOCX with the given content."""
    doc = DocxDocument()

    if headings:
        for heading in headings:
            doc.add_heading(heading, level=1)
            # Add a paragraph under each heading
            idx = headings.index(heading)
            if idx < len(paragraphs):
                doc.add_paragraph(paragraphs[idx])
    else:
        for para in paragraphs:
            doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxProcessor:
    """Tests for DocxProcessor."""

    def setup_method(self) -> None:
        self.processor = DocxProcessor()

    def test_supported_types(self) -> None:
        assert "docx" in self.processor.supported_types
        assert ".docx" in self.processor.supported_types
        assert (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            in self.processor.supported_types
        )

    def test_can_process_docx(self) -> None:
        assert self.processor.can_process("docx") is True
        assert self.processor.can_process(".docx") is True
        assert self.processor.can_process("pdf") is False

    @pytest.mark.asyncio
    async def test_extract_text(self) -> None:
        content = _create_test_docx(["Hello world", "Second paragraph"])
        text = await self.processor.extract_text(content)
        assert "Hello world" in text
        assert "Second paragraph" in text

    @pytest.mark.asyncio
    async def test_extract_text_empty(self) -> None:
        content = _create_test_docx([])
        text = await self.processor.extract_text(content)
        assert text == ""

    @pytest.mark.asyncio
    async def test_process_basic(self) -> None:
        content = _create_test_docx(["Hello world", "This is a test document."])
        result = await self.processor.process(content)
        assert len(result.chunks) >= 1
        assert result.metadata["format"] == "docx"
        assert result.total_characters > 0
        assert result.page_count is None  # DOCX has no fixed pages

    @pytest.mark.asyncio
    async def test_process_with_headings(self) -> None:
        content = _create_test_docx(
            paragraphs=["Introduction text here.", "Details about the topic."],
            headings=["Introduction", "Details"],
        )
        result = await self.processor.process(content)
        assert len(result.chunks) >= 1
        # Check that heading metadata is preserved
        found_heading = False
        for chunk in result.chunks:
            if chunk.metadata.get("heading"):
                found_heading = True
                break
        assert found_heading

    @pytest.mark.asyncio
    async def test_process_long_document(self) -> None:
        """Test that long documents get split into multiple chunks."""
        long_text = "This is a sentence for testing. " * 100
        content = _create_test_docx([long_text])
        result = await self.processor.process(content, chunk_size=500, chunk_overlap=100)
        assert len(result.chunks) > 1

    @pytest.mark.asyncio
    async def test_chunk_indices_are_sequential(self) -> None:
        content = _create_test_docx(["A " * 500, "B " * 500, "C " * 500])
        result = await self.processor.process(content, chunk_size=200, chunk_overlap=50)
        for i, chunk in enumerate(result.chunks):
            assert chunk.chunk_index == i
